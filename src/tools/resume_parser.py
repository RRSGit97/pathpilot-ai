"""
src/tools/resume_parser.py
--------------------------
Parse résumés from three source types and return a structured
:class:`~src.data.schemas.ResumeParseResult`.

Supported formats
-----------------
- **Plain text** – direct passthrough with normalisation.
- **PDF** – page-by-page extraction via ``pypdf`` (no OCR).
- **DOCX** – paragraph extraction via ``python-docx``.

Conservative design rules
-------------------------
- If extracted text is shorter than ``MIN_USEFUL_CHARS`` characters the result
  carries a ``parse_warning`` and ``extracted_skills`` is left empty.  The
  caller must surface this to the user rather than silently proceeding.
- No content is fabricated.  Skills are extracted by a simple keyword scan
  (see :func:`_extract_skills_from_text`); the LLM does the real extraction
  downstream.
- OCR is explicitly not supported in this MVP.  Scanned PDFs will produce an
  empty text extraction and trigger the warning path.

Public API
----------
.. code-block:: python

    from src.tools.resume_parser import parse_resume, ResumeSourceType

    result = parse_resume(path_or_text, source_type=ResumeSourceType.PDF)
    if result.parse_warning:
        st.warning(result.parse_warning)
"""

from __future__ import annotations

import io
import logging
from enum import Enum
from pathlib import Path
from typing import Union

from src.data.schemas import ResumeParseResult
from src.tools.text_utils import (
    extract_keyword_candidates,
    normalize_skill_token,
    normalize_whitespace,
)

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

#: Résumés with fewer extracted characters are flagged as weak / possibly
#: scanned images.  The threshold is intentionally low — most real résumés
#: are 2 000–8 000 characters.
MIN_USEFUL_CHARS: int = 150

#: If a single PDF page produces fewer than this many characters, it is logged
#: but still included (it may be a header page).
MIN_CHARS_PER_PAGE: int = 20


class ResumeSourceType(str, Enum):
    """Identifies how the résumé content is being provided."""

    PLAIN_TEXT = "plain_text"
    PDF = "pdf"
    DOCX = "docx"


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _check_pypdf() -> None:
    """Raise ImportError with a helpful message if pypdf is missing."""
    try:
        import pypdf  # noqa: F401
    except ImportError as exc:
        raise ImportError(
            "pypdf is required for PDF parsing.  Install it with:\n"
            "    pip install pypdf"
        ) from exc


def _check_docx() -> None:
    """Raise ImportError with a helpful message if python-docx is missing."""
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX parsing.  Install it with:\n"
            "    pip install python-docx"
        ) from exc


def _extract_text_from_pdf(source: Union[str, Path, bytes]) -> tuple[str, str | None]:
    """
    Extract plain text from a PDF using pypdf.

    Returns
    -------
    (extracted_text, warning)
        *warning* is ``None`` when extraction looks healthy; a user-facing
        message string when text appears thin or absent (e.g. scanned PDF).
    """
    _check_pypdf()
    import pypdf

    if isinstance(source, bytes):
        reader = pypdf.PdfReader(io.BytesIO(source))
    else:
        reader = pypdf.PdfReader(str(source))

    pages_text: list[str] = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if len(page_text.strip()) < MIN_CHARS_PER_PAGE:
            logger.debug("PDF page %d yielded fewer than %d chars.", i + 1, MIN_CHARS_PER_PAGE)
        pages_text.append(page_text)

    full_text = "\n\n".join(pages_text)
    clean_text = normalize_whitespace(full_text).strip()

    warning: str | None = None
    if len(clean_text) < MIN_USEFUL_CHARS:
        warning = (
            f"PDF text extraction produced only {len(clean_text)} characters — "
            "the file may be a scanned image or have copy-protection enabled. "
            "No OCR is performed in this version. "
            "Please paste your résumé text manually instead."
        )

    return clean_text, warning


def _extract_text_from_docx(source: Union[str, Path, bytes]) -> tuple[str, str | None]:
    """
    Extract plain text from a DOCX file using python-docx.

    Returns
    -------
    (extracted_text, warning)
    """
    _check_docx()
    import docx

    if isinstance(source, bytes):
        doc = docx.Document(io.BytesIO(source))
    else:
        doc = docx.Document(str(source))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    clean_text = normalize_whitespace(full_text).strip()

    warning: str | None = None
    if len(clean_text) < MIN_USEFUL_CHARS:
        warning = (
            f"DOCX extraction produced only {len(clean_text)} characters — "
            "the file may be empty or contain only images/tables. "
            "Please paste your résumé text manually instead."
        )

    return clean_text, warning


def _extract_skills_from_text(text: str) -> list[str]:
    """
    Return a lightweight list of candidate skill tokens found in *text*.

    This is a first-pass heuristic for pre-populating the profile form.
    The LLM extraction node will do the authoritative job — this just
    gives the user something to react to immediately.
    """
    candidates = extract_keyword_candidates(text)
    # Normalise and de-duplicate
    seen: set[str] = set()
    skills: list[str] = []
    for c in candidates:
        norm = normalize_skill_token(c)
        if norm and norm not in seen:
            seen.add(norm)
            skills.append(norm)
    return skills


def _summarise_experience(text: str) -> str:
    """
    Pull the first substantive block of text as a rough experience summary.

    Heuristic: find the first paragraph that is longer than 60 characters
    and return it, trimmed to 500 characters.  This is a placeholder until
    the LLM extraction node runs.
    """
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    for para in paragraphs:
        if len(para) > 60:
            return para[:500] + ("…" if len(para) > 500 else "")
    return ""


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_resume(
    source: Union[str, Path, bytes],
    source_type: ResumeSourceType = ResumeSourceType.PLAIN_TEXT,
) -> ResumeParseResult:
    """
    Parse a résumé from various input types and return structured output.

    Parameters
    ----------
    source:
        - ``str`` → treated as plain text (or file path when *source_type* is
          PDF/DOCX).
        - ``Path`` → file path, used with PDF and DOCX types.
        - ``bytes`` → raw file bytes, used with PDF and DOCX types.
    source_type:
        One of :class:`ResumeSourceType` values.

    Returns
    -------
    ResumeParseResult
        Always returns a result object.  If parsing produced weak output,
        ``parse_warning`` is set and ``extracted_skills`` is empty — the caller
        must not proceed as if the résumé was parsed successfully.

    Raises
    ------
    ValueError
        If *source* is a ``bytes``/``Path`` object but *source_type* is
        ``PLAIN_TEXT`` (ambiguous intent).
    ImportError
        If the required library for the chosen format is not installed.

    Examples
    --------
    >>> result = parse_resume("Alice is a Python developer.", ResumeSourceType.PLAIN_TEXT)
    >>> "python" in result.extracted_skills
    True
    """
    warning: str | None = None

    # --- Route to the right extractor ---
    if source_type == ResumeSourceType.PDF:
        if isinstance(source, str) and not Path(source).exists():
            # Treat as raw text accidentally routed here
            logger.warning("PDF source_type given but source looks like text; treating as plain text.")
            source_type = ResumeSourceType.PLAIN_TEXT
            raw_text = source
        else:
            raw_text, warning = _extract_text_from_pdf(source)

    elif source_type == ResumeSourceType.DOCX:
        if isinstance(source, bytes) or (isinstance(source, (str, Path)) and Path(source).exists()):
            raw_text, warning = _extract_text_from_docx(source)
        else:
            raise ValueError(
                "DOCX source_type requires a file path (str/Path) or raw bytes."
            )

    else:  # PLAIN_TEXT
        if isinstance(source, bytes):
            raise ValueError(
                "source_type is PLAIN_TEXT but bytes were provided.  "
                "Decode first or pass source_type=PDF/DOCX."
            )
        raw_text = str(source)
        raw_text = normalize_whitespace(raw_text).strip()
        if len(raw_text) < MIN_USEFUL_CHARS:
            warning = (
                f"Plain text input is very short ({len(raw_text)} chars).  "
                "Skill extraction may be incomplete."
            )

    # --- Build structured result ---
    if warning:
        # Do not attempt to extract skills from near-empty text
        return ResumeParseResult(
            extracted_skills=[],
            work_experience_summary="",
            education_summary="",
            raw_text=raw_text,
            parse_warning=warning,
        )

    extracted_skills = _extract_skills_from_text(raw_text)
    experience_summary = _summarise_experience(raw_text)

    return ResumeParseResult(
        extracted_skills=extracted_skills,
        work_experience_summary=experience_summary,
        education_summary="",   # LLM node will fill this in
        raw_text=raw_text,
        parse_warning=None,
    )
